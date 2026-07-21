from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/ivy/Documents/Codex/2026-06-17/web")
SOURCE = ROOT / "work/script-export/extracted-game-script.json"
OUT = ROOT / "outputs/替你活到明天-小说版初稿.docx"


CHAPTER_TITLES = {
    "00 序章": ("序章", "雨夜里的旧手机"),
    "01 第一章": ("第一章", "被留下的房间"),
    "02 第二章": ("第二章", "回到高二那年"),
    "03 第三章": ("第三章", "别再往回走"),
    "04 第四章": ("第四章", "校庆前一天"),
    "05 第五章": ("第五章", "证据与玫瑰"),
    "06 第六章": ("第六章", "早一点来"),
    "99 结局": ("尾声", "明天见"),
}


# The novel is a single readable main route. These nodes are player-failure
# branches, duplicate route variants, or "wrong response" beats that would make
# the prose contradictory if placed in the same continuous manuscript.
EXCLUDE_IDS = {
    # First chapter duplicate route without Zhou Xu.
    "ch1_start",
    "ch1_start_02",
    "ch1_living_01",
    "ch1_living_02",
    "ch1_living_03",
    # Second chapter alternate/less intimate responses.
    "ch2_resp_b_01",
    "ch2_resp_c_01",
    "ch2_ignore_01",
    # Third chapter wrong evidence handling.
    "ch3_delete_01",
    # Fourth chapter public confrontation branch.
    "ch4_wall_public_01",
    # Fifth chapter branch kept out of the main emotional route.
    "ch5_envelope_later_01",
    "ch5_call_angry_01",
    # Sixth chapter wrong-route failures.
    "ch6_wrong_zhang_01",
    "ch6_wrong_roof_01",
    # Non-main ending.
    "ch6_fail_end_01",
}


OVERRIDES = {
    "pro_choice": "雨还在下。手机屏幕暗了一次，又被我按亮。我不能再等到明天。先问白祁，再问周叙，然后去陆家。",
    "pro_call_zhou_01": "周叙接了，沉默了两秒，只叫了一声我的名字。我们已经很久没有在深夜打过电话了。分手以后，我们把所有能避免的时间点都避开了，深夜尤其是其中之一。",
    "ch1_choose_search": "房间里还有很多东西没看完。我从书桌开始，把每一个被她留下的角落都看了一遍。",
    "ch1_choose_answer": "我合上抽屉，终于问起那个所有人都绕开的词：高中。",
    "ch1_ans_c_01": "我没有急着把所有发现说出来，只往旁边让了一点。白祁坐下来，没走，也没再问。",
    "ch2_choose_respond": "十七岁的陆眠站在我面前，是真的在等我回答。我知道自己回来了，可她不知道。最后我只能把震惊压下去，跟着她往教室走。",
    "ch2_arrive_02": "有人从我身边跑过去，校服，马尾。然后那个人停住，转过来。是陆眠。她比我记忆里的陆眠小了很多，脸还是那张脸，但眼睛里有一种后来再没见过的东西：干净，轻，没有设防。",
    "ch2_arrive_03": "陆眠看着我，愣了一秒，然后皱起眉：你怎么在这里站着发呆，下午还有课。她认识我。在这里，我们是同学。我是刚转来的学生，她是第一个主动跟我说话的人。",
    "ch2_resp_a_01": "陆眠歪了歪头：想什么？我顿了一下，说：在想你这个人。陆眠愣住，随即笑了，是那种有点不好意思但还是笑出来的表情：转学生都这么奇怪的吗。走啦，一起去教室。",
    "ch2_classroom_01": "我坐在陆眠后排靠窗的位置。这是我在这个时间点的第一天，大家对我还陌生。教室里粉笔灰味很重，窗帘被风吹起一点，窗外有两个男生站在走廊说话。",
    "ch2_classroom_02": "我认出了其中一个。白祁，更小，还带着一点少年气，但眉眼已经是那个样子。另一个我没见过，高，侧脸，站在走廊窗边逆光，看不清表情。陆眠回头，压低声音说：别看那边。",
    "ch2_classroom_03": "我问：为什么？陆眠低下头，声音更轻：别问了。她没有再多说。可她捏紧了校服下摆。我顺着她的目光看过去，只记住走廊窗边那个逆光的侧影。",
    "ch2_sports_02": "你还挺会照顾人。陆眠没躲，只是低头看着我替她别号码布。风吹过来，她忽然笑了一下：那等会儿我跑完，你也要在终点接我。",
    "ch2_choose_after": "走廊很快恢复吵闹，像刚才什么都没发生。可我知道不能当没看见。",
    "ch3_choice_first": "我不能一个人查下去。第一个被我拉进来的人，会决定我接下来得到的是保护、证词，还是另一层沉默。",
    "ch3_choice_threat": "我没有急着回那条短信。比起逞强，我更需要知道对方想让我害怕什么。",
    "ch4_choice_after_note": "这一次，我没有替她决定先去哪儿。我只陪她往她愿意走的地方去。",
    "ch4_choice_wall": "我没有把事情推到人群中央。陆眠最怕被围观，所以我先去切断那只准备按下发送键的手。",
    "ch5_choice_envelope": "信封被拆开前，我先看向白祁。这一次，我不想再让任何人被排除在真相之外。",
    "ch5_choice_call": "电话接通后，我把愤怒压回去。我要的不是吵赢他，而是让他的声音留在证据里。",
    "ch6_final_choice": "最后，我没有公开所有证据，也没有抓住她的手腕。陆眠不需要被我拖回来，她需要有人站到她旁边。",
    "ch2_follow_lu_02": "你不用管这件事。我说：我没说我要管。运动会的时候，你让我站你旁边。现在我也站这儿。陆眠看着镜子里的我，过了很久，才很轻地笑了一下：你真的很奇怪。",
    "ch6_hairtie_02": "它记得你在终点笑，记得你把早饭分一半给我，记得你说要是有人早一点来就好了。我看着她：陆眠，我来了。",
}


SECTION_OPENERS = {
    "00 序章": "那天晚上，我第一次明白，遗物也会敲门。",
    "01 第一章": "陆家的灯太安静了。安静到我一推开门，就像踩进一场还没结束的葬礼。",
    "02 第二章": "旧手机亮起来的时候，我以为自己只是太累。直到雨声消失，上课铃从十年前的操场尽头响起。",
    "03 第三章": "从过去醒来以后，现实没有变轻。它只是把更冷的一面翻给我看。",
    "04 第四章": "第二次回到高二，我终于知道，救一个人不能只靠知道答案。",
    "05 第五章": "现实里的证据一块块浮上来，像水底的玻璃。每一片都反光，也都割手。",
    "06 第六章": "我回到坠落前二十分钟。所有后来才懂的事，都必须在这二十分钟里变成一句她听得进去的话。",
    "99 结局": "天亮的时候，我没有再听见那场雨。",
}


def set_run_font(run, name: str = "Songti SC", size: float | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def replace_player_pronouns(text: str) -> str:
    replacements = [
        ("你们", "我们"),
        ("你的", "我的"),
        ("你自己", "我自己"),
        ("你也", "我也"),
        ("你还", "我还"),
        ("你已经", "我已经"),
        ("你终于", "我终于"),
        ("你第一次", "我第一次"),
        ("你感觉", "我感觉"),
        ("你打开", "我打开"),
        ("你打下", "我打下"),
        ("你坐在", "我坐在"),
        ("你站在", "我站在"),
        ("你坐到", "我坐到"),
        ("你认出", "我认出"),
        ("你知道", "我知道"),
        ("你问", "我问"),
        ("你说", "我说"),
        ("你听见", "我听见"),
        ("你看见", "我看见"),
        ("你没有", "我没有"),
        ("你只", "我只"),
        ("你不", "我不"),
        ("你把", "我把"),
        ("你摸到", "我摸到"),
        ("你记住", "我记住"),
        ("你来", "我来"),
        ("你跟", "我跟"),
        ("看着你", "看着我"),
        ("等你", "等我"),
        ("发现你", "发现我"),
        ("在你", "在我"),
        ("跟你", "跟我"),
        ("向你", "向我"),
        ("给你", "给我"),
        ("陪你", "陪我"),
        ("替你", "替我"),
    ]
    for src, dst in replacements:
        text = text.replace(src, dst)
    return text


def normalize_text(node: dict, zhang_revealed: bool) -> tuple[str, bool]:
    node_id = node["id"]
    overridden = node_id in OVERRIDES
    text = OVERRIDES.get(node_id)
    if not text:
        # Before the heroine can identify Zhang Heng, use the masked runtime
        # display text. At the darkroom reveal and afterward, allow the name.
        if zhang_revealed or node_id in {"ch4_photo_room_02", "ch4_photo_room_04"}:
            text = node.get("sourceText") or node.get("defaultDisplayText") or ""
        else:
            text = node.get("defaultDisplayText") or node.get("sourceText") or ""

    if node_id == "ch4_photo_room_02":
        zhang_revealed = True

    text = text.strip()
    if not text:
        return "", zhang_revealed

    # Convert player-addressed narration into first-person prose while trying
    # not to rewrite quoted speech after a Chinese colon.
    if overridden:
        pass
    elif "：" in text:
        head, tail = text.split("：", 1)
        text = f"{replace_player_pronouns(head)}：{tail}"
    elif node.get("speaker") in {"沈栀", "旁白"} or any(name in text for name in ["陆眠", "白祁", "周叙", "那个人", "张恒"]):
        text = replace_player_pronouns(text)

    text = text.replace("你打开短信输入框", "我打开短信输入框")
    text = text.replace("你打下一句话", "我打下一句话")
    text = text.replace("你感觉自己", "我感觉自己")
    text = text.replace("你站在高中操场边", "我站在高中操场边")
    text = text.replace("你手里拿着", "我手里拿着")
    text = text.replace("你坐在陆眠后排", "我坐在陆眠后排")
    text = text.replace("你们已经很久", "我们已经很久")
    text = text.replace("你听见那边", "我听见那边")
    text = text.replace("你说：", "我说：")

    # Remove a few remaining game-UI traces.
    text = text.replace("现在轮到我问了。", "")
    text = text.replace("我可以问高中，也可以先问他有没有好好睡觉。", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip(), zhang_revealed


def is_body_node(node: dict) -> bool:
    group = node.get("group", "")
    if group not in CHAPTER_TITLES:
        return False
    if node["id"] in EXCLUDE_IDS:
        return False
    if node.get("ending") == "fail":
        return False
    if "任务失败" in (node.get("chapter") or "") or "坏结局" in (node.get("chapter") or ""):
        return False
    return True


def add_paragraph(doc: Document, text: str, style: str = "Normal", first_line: bool = True):
    p = doc.add_paragraph(style=style)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.45
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_heading(doc: Document, label: str, title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    set_run_font(r, "Songti SC", 15, True)
    r.font.color.rgb = RGBColor(74, 66, 54)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(16)
    r2 = p2.add_run(title)
    set_run_font(r2, "Songti SC", 20, True)
    r2.font.color.rgb = RGBColor(34, 34, 34)


def build_docx():
    data = json.loads(SOURCE.read_text(encoding="utf-8"))
    nodes = [n for n in data["nodes"] if is_body_node(n)]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Songti SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    title.paragraph_format.space_after = Pt(12)
    r = title.add_run("替你活到明天")
    set_run_font(r, "Songti SC", 28, True)
    r.font.color.rgb = RGBColor(31, 31, 31)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(70)
    r = subtitle.add_run("小说版初稿")
    set_run_font(r, "Songti SC", 15, False)
    r.font.color.rgb = RGBColor(110, 96, 80)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(6)
    r = note.add_run("基于当前游戏主线改写 · 沈栀第一人称 · 限制视角")
    set_run_font(r, "Songti SC", 10, False)
    r.font.color.rgb = RGBColor(120, 120, 120)

    stamp = doc.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = stamp.add_run(f"生成日期：{datetime.now().strftime('%Y-%m-%d')}")
    set_run_font(r, "Songti SC", 9, False)
    r.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    current_group = None
    zhang_revealed = False
    for node in nodes:
        group = node["group"]
        if group != current_group:
            if current_group is not None:
                doc.add_section(WD_SECTION.NEW_PAGE)
            current_group = group
            label, title_text = CHAPTER_TITLES[group]
            add_heading(doc, label, title_text)
            opener = SECTION_OPENERS.get(group)
            if opener:
                p = add_paragraph(doc, opener, first_line=False)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].italic = True
                p.runs[0].font.color.rgb = RGBColor(122, 111, 96)

        text, zhang_revealed = normalize_text(node, zhang_revealed)
        if not text:
            continue

        # Small scene breaths at major turns, not every node, to avoid the
        # "game script" feeling while keeping the reader oriented.
        if node["id"] in {
            "pro_search_01",
            "ch1_room_01",
            "ch2_arrive_01",
            "ch2_sports_01",
            "ch3_school_01",
            "ch4_artroom_01",
            "ch5_studio_01",
            "ch6_broadcast_01",
        }:
            breath = doc.add_paragraph()
            breath.alignment = WD_ALIGN_PARAGRAPH.CENTER
            breath.paragraph_format.space_before = Pt(8)
            breath.paragraph_format.space_after = Pt(8)
            rr = breath.add_run("*")
            set_run_font(rr, "Songti SC", 12, False)
            rr.font.color.rgb = RGBColor(160, 150, 135)

        add_paragraph(doc, text)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("《替你活到明天》小说版初稿")
    set_run_font(fr, "Songti SC", 8, False)
    fr.font.color.rgb = RGBColor(140, 140, 140)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    out = build_docx()
    print(out)
